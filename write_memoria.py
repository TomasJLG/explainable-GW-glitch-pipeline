"""
write_memoria.py  — versión compacta (≤50 páginas)
Genera Memoria_PI_LIGO.docx
Autor: Tomás Legal · Compañeros: Daniel Tanco, Miguel Peralias
"""
import json
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
FIG  = ROOT / "figures"
FIG.mkdir(exist_ok=True)

# ─── XML helpers ──────────────────────────────────────────────────────────────

def _pPr(obj):
    """Return the pPr element of a paragraph, style element, or raw pPr."""
    if hasattr(obj, '_p'):           return obj._p.get_or_add_pPr()
    if hasattr(obj, 'get_or_add_pPr'): return obj
    return obj  # already a pPr element

def _spacing(obj, lines=360, bef=0, aft=6):
    pPr = _pPr(obj)
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    s = OxmlElement('w:spacing')
    s.set(qn('w:line'),    str(lines))
    s.set(qn('w:lineRule'),'auto')
    s.set(qn('w:before'),  str(int(bef*20)))
    s.set(qn('w:after'),   str(int(aft*20)))
    pPr.append(s)

def _shd(para, fill='F2F2F2'):
    pPr = para._p.get_or_add_pPr()
    e = OxmlElement('w:shd')
    e.set(qn('w:val'),   'clear')
    e.set(qn('w:color'), 'auto')
    e.set(qn('w:fill'),  fill)
    pPr.append(e)

# ─── Document setup ────────────────────────────────────────────────────────────

def setup():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(3.0)

    def sty(name, size, bold=False, italic=False,
            rgb=None, lt=360, bef=0, aft=6):
        st = doc.styles[name]
        st.font.name  = 'Calibri'
        st.font.size  = Pt(size)
        st.font.bold  = bold
        st.font.italic= italic
        if rgb:
            st.font.color.rgb = RGBColor(*rgb)
        _spacing(st._element.get_or_add_pPr(), lt, bef, aft)

    sty('Normal',    13, lt=360, bef=0, aft=5)
    sty('Heading 1', 15, bold=True,  rgb=(0x1F,0x49,0x7D), lt=360, bef=14, aft=6)
    sty('Heading 2', 13, bold=True,  rgb=(0x2E,0x74,0xB5), lt=360, bef=10, aft=4)
    sty('Heading 3', 12, bold=True,  italic=True,
        rgb=(0x44,0x44,0x44), lt=360, bef=6, aft=3)
    return doc

# ─── Para / heading helpers ────────────────────────────────────────────────────

def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def p(doc, text, bold=False, italic=False, center=False, aft=5):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name   = 'Calibri'
    run.font.size   = Pt(13)
    run.font.bold   = bold
    run.font.italic = italic
    para.alignment  = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(para, 360, 0, aft)
    return para

def bullet(doc, label, text, bld=True):
    para = doc.add_paragraph()
    r1 = para.add_run(f'• {label}: ')
    r1.font.name = 'Calibri'; r1.font.size = Pt(13); r1.font.bold = bld
    r2 = para.add_run(text)
    r2.font.name = 'Calibri'; r2.font.size = Pt(13)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(para, 360, 0, 4)

def fig(doc, path, caption, n):
    path = Path(path)
    if not path.exists():
        p(doc, f'[Figura {n} no disponible: {path.name}]', italic=True, center=True)
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(para, 360, 4, 2)
    para.add_run().add_picture(str(path), width=Cm(13))
    cap = doc.add_paragraph()
    r = cap.add_run(f'Figura {n}. {caption}')
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.italic = True
    r.font.color.rgb = RGBColor(0x44,0x44,0x44)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(cap, 360, 0, 8)

def tbl(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i, hd in enumerate(headers):
        c = hr.cells[i]
        c.text = hd
        r = c.paragraphs[0].runs[0]
        r.font.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'),'2E74B5')
        c._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.name = 'Calibri'
            c.paragraphs[0].runs[0].font.size = Pt(11)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for ci, w in enumerate(widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()

def code_file(doc, label, fpath, purpose):
    h(doc, f'{label}  {Path(fpath).name}', level=3)
    p(doc, f'Propósito: {purpose}', italic=True, aft=3)
    path = Path(fpath)
    if not path.exists():
        p(doc, f'[Archivo no encontrado: {fpath}]', italic=True); return
    with open(str(path), encoding='utf-8', errors='replace') as f:
        txt = f.read()
    if path.suffix in ('.json', '.ipynb'):
        try:
            nb = json.loads(txt)
            cells = nb.get('cells', [])
            lines = []
            for cell in cells:
                if cell.get('cell_type') == 'code':
                    src = cell.get('source', [])
                    if isinstance(src, list): src = ''.join(src)
                    if src.strip():
                        lines.append('# ── celda ──────────────────────')
                        lines.append(src)
            txt = '\n'.join(lines) if lines else txt
        except Exception:
            pass
    for line in txt.split('\n'):
        para = doc.add_paragraph()
        run  = para.add_run(line if line else ' ')
        run.font.name = 'Consolas'; run.font.size = Pt(9)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        _shd(para)
        pPr = para._p.get_or_add_pPr()
        s = OxmlElement('w:spacing')
        s.set(qn('w:line'),'200'); s.set(qn('w:lineRule'),'auto')
        s.set(qn('w:before'),'0'); s.set(qn('w:after'),'0')
        pPr.append(s)
    doc.add_paragraph()

def page_numbers(doc):
    for sec in doc.sections:
        footer = sec.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()
        run = fp.add_run()
        run.font.name = 'Calibri'; run.font.size = Pt(11)
        for tag, txt in [('begin',''), ('','PAGE'), ('end','')]:
            if tag:
                e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), tag)
                run._r.append(e)
            else:
                e = OxmlElement('w:instrText'); e.text = txt
                run._r.append(e)

# ─── Architecture diagram ──────────────────────────────────────────────────────

def gen_arch():
    fig_obj, ax = plt.subplots(figsize=(14, 4))
    ax.set_xlim(0, 14); ax.set_ylim(0, 4); ax.axis('off')
    fig_obj.patch.set_facecolor('white')
    COLORS = ['#D5E8D4','#DAE8FC','#FFF2CC','#DAE8FC','#FFF2CC','#E1D5E7','#FFE6CC']
    boxes = [
        (0.1, 1.2, 1.6, 1.6, 'GWOSC\nstrain HDF5', '4096 Hz'),
        (2.0, 0.9, 2.2, 2.2, 'M0\nGenerador', 'Q-transform\n128×128'),
        (4.5, 1.2, 1.6, 1.6, 'NPZ\nDatasets', '6465+1700'),
        (6.4, 0.9, 2.4, 2.2, 'M1\nGlitchAE', 'latent_dim=32\nMSE nominal'),
        (9.1, 1.2, 1.6, 1.6, 'Features\n6D', 'ae·log_e·snr\npfreq·bw·dur'),
        (11.0, 0.9, 2.0, 2.2, 'M2\nANFIS v3', 'Stage1+Stage2\nnorm IFO×época'),
        (13.2, 1.2, 0.7, 1.6, 'Clase\nmacro', '×4'),
    ]
    for i, (x, y, w, hh, lbl, sub) in enumerate(boxes):
        rect = FancyBboxPatch((x,y),w,hh, boxstyle='round,pad=0.07',
                              facecolor=COLORS[i], edgecolor='#555', linewidth=1.1)
        ax.add_patch(rect)
        ax.text(x+w/2, y+hh*0.65, lbl, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color='#1a1a2e')
        ax.text(x+w/2, y+hh*0.25, sub, ha='center', va='center',
                fontsize=7, color='#555', style='italic')
    arrows = [(1.7,2.0),(4.2,4.5),(6.1,6.4),(8.8,9.1),(10.7,11.0),(13.0,13.2)]
    for x1, x2 in arrows:
        ax.annotate('', xy=(x2,2.0), xytext=(x1,2.0),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.4))
    ax.set_title('Pipeline M0 → M1 → M2 para detección y clasificación de glitches LIGO O3',
                 fontsize=10, fontweight='bold', color='#1a1a2e', pad=10)
    out = FIG / 'architecture_pipeline.png'
    fig_obj.savefig(str(out), dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig_obj)
    return out

# ─── Portada ──────────────────────────────────────────────────────────────────

def cover(doc):
    for _ in range(4): doc.add_paragraph()
    t = doc.add_paragraph()
    r = t.add_run('Pipeline explicable para la detección y\nclasificación de glitches en datos LIGO O3')
    r.font.name = 'Calibri'; r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(t, 360, 0, 10)
    s = doc.add_paragraph()
    rs = s.add_run('Proyecto Intermodular — Ciclo Formativo de Grado Superior en IA y Big Data')
    rs.font.name = 'Calibri'; rs.font.size = Pt(14); rs.font.italic = True
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(s, 360, 0, 30)
    for _ in range(3): doc.add_paragraph()
    for lbl, val in [
        ('Alumno principal:',      'Tomás Legal'),
        ('Compañeros de proyecto:','Daniel Tanco · Miguel Peralias'),
        ('Curso:',                 '2025–2026'),
        ('Fecha de entrega:',      'Mayo de 2026'),
    ]:
        ln = doc.add_paragraph()
        r1 = ln.add_run(f'{lbl}  '); r1.font.name='Calibri'; r1.font.size=Pt(13); r1.font.bold=True
        r2 = ln.add_run(val);        r2.font.name='Calibri'; r2.font.size=Pt(13)
        ln.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(ln, 360, 0, 6)
    doc.add_page_break()

# ─── Capítulo 1 ───────────────────────────────────────────────────────────────

def cap1(doc):
    h(doc, '1. Introducción')

    h(doc, '1.1 Contexto y justificación', 2)
    p(doc,
      'Los interferómetros LIGO (Laser Interferometer Gravitational-Wave Observatory) '
      'detectaron la primera señal de ondas gravitacionales en 2015 [1], inaugurando '
      'una nueva ventana de observación del universo. Durante la tercera campaña O3 '
      '(2019–2020), los detectores H1 (Hanford) y L1 (Livingston) operaron con '
      'sensibilidad de 10⁻²³ m/√Hz, lo que los hace susceptibles a transitorios de '
      'ruido instrumental denominados glitches [2]. El proyecto Gravity Spy [3] clasifica '
      '23 morfologías de glitch mediante CNNs y ciencia ciudadana, pero sus modelos '
      'no ofrecen explicaciones lingüísticas de las decisiones, limitando el diagnóstico '
      'instrumental. Este proyecto desarrolla un pipeline de tres módulos con un '
      'clasificador difuso explícito para cubrir esa brecha.')

    h(doc, '1.2 Definición del problema', 2)
    p(doc,
      'Un glitch es un exceso de potencia transitorio (duración < 10 s) visible en la '
      'representación tiempo-frecuencia Q-transform [5] de los datos de tensión. '
      'El problema consiste en: (1) detectar si una ventana Q-transform de 128×128 '
      'píxeles contiene un glitch (M1, detección no supervisada) y (2) clasificarlo '
      'en una de las 4 macro-clases Loud, Burst, Scatter u Other (M2, clasificación '
      'supervisada). Las 23 clases finas de Gravity Spy se agrupan en estas 4 '
      'macro-clases según similitud morfológica y física.')

    h(doc, '1.3 Objetivo general', 2)
    p(doc,
      'Diseñar, implementar y validar un pipeline explicable de tres módulos '
      '(M0 — generador Q-transform, M1 — autoencoder de detección, '
      'M2 — clasificador ANFIS) para la detección y clasificación de glitches '
      'en datos LIGO O3, garantizando reproducibilidad e interpretabilidad '
      'de las decisiones de clasificación.')

    h(doc, '1.4 Objetivos específicos', 2)
    for code, txt in [
        ('OE1', 'Implementar M0 como generador NPZ reanudable con checkpoints para 4 '
                'combinaciones IFO×época (H1/O3a, H1/O3b, L1/O3a, L1/O3b).'),
        ('OE2', 'Entrenar GlitchAE (M1) con latent_dim=32 sobre datos nominales, '
                'alcanzando AUROC > 0,80 sobre run03 etiquetado.'),
        ('OE3', 'Extraer un vector de 6 características por ventana combinando '
                'el score AE con metadatos físicos del trigger (SNR, frecuencia, '
                'ancho de banda, duración, log-energía).'),
        ('OE4', 'Entrenar ANFIS jerárquico (M2 v3) con normalización por IFO×época, '
                'logrando accuracy > 0,65 y macro-F1 > 0,60.'),
        ('OE5', 'Validar el pipeline E2E sobre 500 muestras etiquetadas con '
                'consistencia > 95 % y latencia < 50 ms/muestra.'),
        ('OE6', 'Documentar las decisiones de diseño clave mediante ADRs en el código.'),
    ]:
        bullet(doc, code, txt)

# ─── Capítulo 2 ───────────────────────────────────────────────────────────────

def cap2(doc):
    h(doc, '2. Marco Teórico')

    h(doc, '2.1 LIGO y los glitches gravitacionales', 2)
    p(doc,
      'LIGO opera como interferómetro de Michelson con brazos de 4 km. Las ondas '
      'gravitacionales producen cambios diferenciales de longitud de ~10⁻²¹ m [2]. '
      'Los glitches son transitorios instrumentales de origen sísmico, acústico, '
      'electromagnético o de control que ocurren a razón de ~1/min. '
      'Su clasificación sistemática (detector characterization) es esencial para '
      'distinguirlos de señales astrofísicas reales [3]. Los datos de tensión están '
      'disponibles públicamente en el LIGO Open Science Center (GWOSC) [4] en '
      'bloques HDF5 de 4096 s a 4096 Hz.')

    h(doc, '2.2 Transformada Q para señales transitorias', 2)
    p(doc,
      'La Q-transform [5] descompone una señal en teselas tiempo-frecuencia de '
      'resolución variable: ventanas largas a bajas frecuencias (alta resolución '
      'espectral) y cortas a altas frecuencias (alta resolución temporal), '
      'manteniendo constante Q = f/Δf. En este proyecto se aplica con '
      'Q ∈ [4, 64], f ∈ [20, 1700] Hz, sobre segmentos de ±8,5 s de contexto, '
      'extrayendo la ventana central de ±0,5 s como imagen 128×128 píxeles. '
      'El preprocesado de gwpy [10] incluye blanqueado (whitening), filtro '
      'pasa-banda y notch en armónicos de la red (60/120/180 Hz).')

    h(doc, '2.3 Autoencoders convolucionales para detección de anomalías', 2)
    p(doc,
      'Un autoencoder [6] se entrena para reconstruir su entrada a través de un '
      'cuello de botella latente. Entrenado exclusivamente sobre ruido nominal, '
      'el error de reconstrucción (MSE) actúa como puntuación de anomalía: '
      'los glitches, al no haber sido vistos durante el entrenamiento, se '
      'reconstruyen con mayor error. GlitchAE implementa 4 bloques Conv2d '
      '(1→32→64→128→256, stride=2), dos capas FC que comprimen a latent_dim=32, '
      'y un decoder simétrico con ConvTranspose2d y Sigmoid. El tamaño latente '
      'reducido (latent_dim=32) es deliberado: fuerza el bottleneck y maximiza '
      'la discriminación nominal/glitch (AUROC sube de 0,53 con dim=128 a 0,85 '
      'con dim=32).')

    h(doc, '2.4 ANFIS Takagi-Sugeno', 2)
    p(doc,
      'ANFIS [7] implementa un sistema de inferencia difusa Takagi-Sugeno de primer '
      'orden como una red de cinco capas diferenciable. Cada regla tiene la forma '
      '"SI x₁ es A₁ Y … ENTONCES y = p₁x₁ + … + r", donde A_i son funciones '
      'de pertenencia GBell: μ(x;a,b,c)=1/(1+|(x−c)/a|^(2b)). El entrenamiento '
      'híbrido [7] alterna LSE (actualiza consecuentes lineales de forma óptima '
      'mediante pseudoinversa) con Adam [9] (actualiza parámetros no lineales de '
      'las MFs). Los centros iniciales de reglas se obtienen por agrupamiento '
      'sustractivo [8], donde el radio ra controla el número de reglas emergentes '
      'sin necesidad de especificarlo a priori.')

    h(doc, '2.5 Estado del arte', 2)
    p(doc,
      'Gravity Spy [3] es el sistema de referencia para clasificación de glitches '
      'LIGO: CNN + ciencia ciudadana, 23 clases, >1 M ventanas etiquetadas en O3. '
      'Vajente et al. [12] proponen sustraer ruido no estacionario con ML. '
      'Los autoencoders variacionales y one-class SVMs se han explorado para '
      'detección no supervisada, pero ningún trabajo previo combina embedding '
      'de autoencoder con ANFIS explícito para clasificación de glitches, '
      'lo que constituye la contribución diferencial de este proyecto.')

# ─── Capítulo 3 ───────────────────────────────────────────────────────────────

def cap3(doc):
    h(doc, '3. Marco Metodológico')

    h(doc, '3.1 Metodología de desarrollo', 2)
    p(doc,
      'Se sigue un modelo iterativo-incremental en tres fases (M0, M1, M2), '
      'cada una con ciclo diseño → implementación → evaluación → ADR. '
      'Las decisiones clave se documentan como registros de decisión arquitectónica: '
      'ADR-001 (raw sin normalización por ventana), ADR-002 (latent_dim=32), '
      'ADR-003 (split temporal P80), ADR-004 (P1/P99 global solo del train), '
      'ADR-005 (4 macro-clases), ADR-006 (ANFIS jerárquico) y '
      'ADR-007 (normalización P2/P98 por IFO×época).')

    h(doc, '3.2 Herramientas', 2)
    tbl(doc,
        headers=['Herramienta', 'Versión', 'Uso en el proyecto'],
        rows=[
            ['Python',     '3.11',   'Lenguaje principal del pipeline'],
            ['PyTorch [13]','2.x',   'GlitchAE: definición, entrenamiento, inferencia GPU'],
            ['gwpy [10]',  '3.x',    'Q-transform, blanqueado, lectura de strain HDF5'],
            ['gwosc',      '0.7',    'Localización y descarga de bloques de strain GWOSC [4]'],
            ['scipy/numpy','≥1.10',  'Agrupamiento sustractivo, PCA (SVD), LSE, métricas'],
            ['scikit-learn','1.x',   'AUROC, AUPRC, TPR@FPR1%, matrices de confusión'],
            ['matplotlib', '3.x',    'Curvas ROC, pérdida, MFs, matrices de confusión'],
            ['python-docx','1.x',    'Generación de este documento'],
        ],
        widths=[3.0, 2.0, 7.5]
    )

    h(doc, '3.3 Proceso de desarrollo', 2)
    p(doc,
      'M0 genera tres datasets NPZ: run02v2 (GPSs aleatorios uniformes en '
      'segmentos de ciencia, sin etiqueta, para train M1), run03 (centrado en '
      'triggers Gravity Spy, SNR≥7,5, H1/O3a, 500 ventanas etiquetadas) y '
      'run03_minority (4 combos × 300 ventanas, clases minoritarias). '
      'M1 entrena GlitchAE sobre run02v2 con split temporal P80 y normalización '
      'P1/P99 del train. M2 extrae el vector de 6 features (ae_score + 5 metadatos '
      'CSV) y entrena el ANFIS jerárquico v3 sobre los 1637 eventos únicos '
      'deduplicados de run03 + run03_minority. La validación E2E usa '
      'm2_pipeline_eval.py sobre las 500 ventanas etiquetadas de run03.')

    h(doc, '3.4 Cronograma', 2)
    tbl(doc,
        headers=['Fase', 'Actividad', 'Período', 'Duración'],
        rows=[
            ['M0','run02v2 (H1/O3a) + diseño generadores','Oct 2025','3 sem.'],
            ['M0','Expansión 4 combos + run03 + run03_minority','Nov–Dic 2025','5 sem.'],
            ['M1','GlitchAE v1/v2 (Kaggle) + v3 run02v2','Ene–Feb 2026','5 sem.'],
            ['M2','Features v1 + ANFIS v2 plano','Mar 2026','3 sem.'],
            ['M2','ANFIS v3 jerárquico + norma IFO×época + E2E','Abr 2026','3 sem.'],
            ['Doc','Memoria PI e informes finales','May 2026','2 sem.'],
        ],
        widths=[1.5, 6.0, 2.5, 2.5]
    )

    h(doc, '3.5 Presupuesto', 2)
    tbl(doc,
        headers=['Concepto', 'Coste', 'Nota'],
        rows=[
            ['Software (Python, PyTorch, gwpy…)', '0 €','Licencias MIT/BSD/Apache'],
            ['Datos GWOSC + Gravity Spy CSV',      '0 €','Acceso público'],
            ['Cómputo GPU (Kaggle)',                '0 €','Cuota gratuita 30 h/sem.'],
            ['Hardware personal (laptop)',          '—',  'Amortizado, no imputable'],
            ['Total imputable al PI',               '0 €','—'],
        ],
        widths=[5.5, 2.0, 5.0]
    )

# ─── Capítulo 4 ───────────────────────────────────────────────────────────────

def cap4(doc):
    h(doc, '4. Resultados')
    n = [1]  # mutable counter

    def nf(): v = n[0]; n[0] += 1; return v

    h(doc, '4.1 Diagrama de arquitectura del pipeline', 2)
    p(doc, f'La Figura {n[0]} muestra el flujo completo M0 → M1 → M2.')
    fig(doc, FIG/'architecture_pipeline.png',
        'Diagrama de bloques del pipeline M0→M1→M2.', nf())

    h(doc, '4.2 Módulo M0: datasets generados', 2)
    tbl(doc,
        headers=['Dataset','IFO','Época','N ventanas','Etiquetado','Uso'],
        rows=[
            ['run02v2','H1','O3a','465','No','Train M1'],
            ['run02v2','H1','O3b','2000','No','Train M1'],
            ['run02v2','L1','O3a','2000','No','Train M1'],
            ['run02v2','L1','O3b','2000','No','Train M1'],
            ['run03','H1','O3a','500','Sí (23 clases)','Eval M1+M2'],
            ['run03_minority','×4 combos','O3a/b','300 c/u','Sí','Train+Eval M2'],
            ['Total etiquetado (dedup.)','—','—','1637','—','M2'],
        ],
        widths=[2.8,1.1,1.1,2.0,2.8,2.6]
    )
    p(doc,
      'La diferencia entre H1/O3a (465 ventanas) y el resto (2000) se debe a la '
      'disponibilidad de segmentos de ciencia en GWOSC para esa combinación. '
      'La distribución de macro-clases en el dataset etiquetado es: '
      'Loud=279 (17 %), Burst=396 (24 %), Scatter=554 (34 %), Other=408 (25 %).')

    h(doc, '4.3 Módulo M1: GlitchAE — entrenamiento y evaluación', 2)
    p(doc,
      'GlitchAE v3 (latent_dim=32, 17,9M parámetros) se entrenó sobre las '
      '~6140 ventanas nominales de run02v2 (filtro P95 log_energy, split P80 temporal). '
      'Hiperparámetros: Adam lr=1e-3, weight_decay=1e-5, batch=64, epochs=80, '
      'patience=15. La pérdida mínima de validación fue best_val_loss=0,04199 '
      '(época 57). Normalización: p1=0,0249, p99=7,354 (calculada solo del train).')
    fig(doc, ROOT/'m1_v3_outputs'/'training_curve_v3.png',
        'Curvas de pérdida MSE train/val de GlitchAE v3. Convergencia estable a partir de la época 8.', nf())
    fig(doc, ROOT/'m1_v3_outputs'/'roc_curve_v3.png',
        'Curva ROC de M1 (GlitchAE v3) sobre run03 — 500 ventanas etiquetadas. AUROC = 0,85.', nf())
    fig(doc, ROOT/'m1_v3_outputs'/'score_by_class_v3.png',
        'Distribución del score de anomalía (MSE) por macro-clase. Loud presenta los scores más elevados.', nf())
    p(doc,
      'AUROC = 0,85 supera el objetivo OE2 (> 0,80) y mejora significativamente '
      'respecto a v1 (0,527) y v2 (0,305). La mejora se debe al uso de run02v2 '
      'con GPS uniformes, al bottleneck latent_dim=32 y a la normalización P1/P99 '
      'aplicada consistentemente.')

    h(doc, '4.4 Módulo M2 v2: ANFIS plano', 2)
    tbl(doc,
        headers=['Clase','F1-score','Precisión','Recall'],
        rows=[
            ['Loud',   '0,800','0,83','0,77'],
            ['Burst',  '0,784','0,78','0,79'],
            ['Scatter','0,728','0,74','0,72'],
            ['Other',  '0,359','0,45','0,28'],
            ['Global (macro-F1 / accuracy)','0,6677 / —','—','0,7460'],
        ],
        widths=[3.5,3.0,3.0,3.0]
    )
    fig(doc, ROOT/'m2_outputs_v2'/'confusion_matrix_v2.png',
        'Matriz de confusión ANFIS v2 (plano). La clase Other presenta la mayor tasa de error.', nf())

    h(doc, '4.5 Módulo M2 v3: ANFIS jerárquico con normalización IFO×época', 2)
    p(doc,
      'La versión v3 estructura el problema en dos etapas: Stage 1 clasifica '
      'Loud vs. Rest (ra=0,15, 7 reglas) y Stage 2 clasifica Burst/Scatter/Other '
      'sobre los ejemplos predichos como Rest (ra=0,10, 8 reglas). '
      'La normalización P2/P98 se ajusta independientemente para cada grupo '
      'H1_O3a, H1_O3b, L1_O3a, L1_O3b sobre el conjunto de train, '
      'reduciendo el sesgo inter-detector e inter-época.')
    tbl(doc,
        headers=['Etapa','Tarea','Reglas','Métrica','Valor'],
        rows=[
            ['Stage 1','Loud vs. Rest','7','Accuracy','0,903'],
            ['Stage 2','Burst/Scatter/Other','8','Macro-F1','0,690'],
            ['Combinado','4 clases','15','Accuracy / Macro-F1','0,694 / 0,6815'],
        ],
        widths=[2.0,4.0,2.0,2.5,2.0]
    )
    fig(doc, ROOT/'m2_outputs_v3'/'confusion_matrix_v3.png',
        'Matriz de confusión ANFIS v3 jerárquico. Accuracy=0,694, Macro-F1=0,6815.', nf())
    fig(doc, ROOT/'m2_outputs_v3'/'membership_functions_stage1.png',
        'Funciones de pertenencia GBell — Stage 1 (Loud vs. Rest, 7 reglas).', nf())
    fig(doc, ROOT/'m2_outputs_v3'/'training_curves_stage1.png',
        'Curvas de pérdida CrossEntropy Stage 1 (aprendizaje híbrido LSE+Adam).', nf())

    h(doc, '4.6 Validación E2E del pipeline', 2)
    tbl(doc,
        headers=['Métrica','Valor'],
        rows=[
            ['Ventanas procesadas','500 / 500'],
            ['Consistencia de predicciones','100 % (500/500)'],
            ['Latencia de inferencia (CPU)','~17 ms/muestra'],
        ],
        widths=[6.0,6.5]
    )
    fig(doc, ROOT/'m2_outputs'/'pipeline_eval'/'pipeline_confusion_matrix.png',
        'Matriz de confusión del pipeline E2E M1→M2 v2 sobre run03 (500 ventanas).', nf())
    fig(doc, ROOT/'m2_outputs'/'pipeline_eval'/'rule_firing_heatmap.png',
        'Heatmap de activación de reglas ANFIS por clase. Permite diagnóstico instrumental explícito.', nf())

# ─── Capítulo 5 ───────────────────────────────────────────────────────────────

def cap5(doc):
    h(doc, '5. Conclusiones')

    h(doc, '5.1 Valoración del trabajo', 2)
    p(doc,
      'El proyecto entrega un pipeline funcional y reproducible que cumple o supera '
      'todos los objetivos específicos fijados: AUROC=0,85 (OE2), vector 6D de '
      'features (OE3), ANFIS v3 con accuracy=0,694 y macro-F1=0,6815 (OE4), '
      'validación E2E 500/500 a 17 ms/muestra (OE5). La adopción del ANFIS '
      'Takagi-Sugeno como clasificador final garantiza que cada predicción puede '
      'expresarse en lenguaje natural, conectando la morfología del glitch '
      'con los parámetros físicos del interferómetro.')

    h(doc, '5.2 Objetivos alcanzados', 2)
    tbl(doc,
        headers=['Obj.','Estado','Evidencia clave'],
        rows=[
            ['OE1','✓','6465 + 1700 ventanas NPZ en 4 combos IFO×época'],
            ['OE2','✓ (superado)','AUROC = 0,85 > 0,80 sobre run03 (500 muestras)'],
            ['OE3','✓','m2_features_v2.npz — 1637 muestras, 6 features'],
            ['OE4','✓ (superado)','Acc=0,694 > 0,65 ; macro-F1=0,6815 > 0,60'],
            ['OE5','✓ (superado)','500/500 consistentes, 17 ms < 50 ms'],
            ['OE6','✓','ADR-001 a ADR-007 en código y memoria'],
        ],
        widths=[1.5,2.5,8.5]
    )

    h(doc, '5.3 Dificultades superadas', 2)
    for lbl, txt in [
        ('Data leakage temporal',
         'Split P80 del tiempo GPS en lugar de aleatorio, evitando que el train '
         'vea ruido posterior al val.'),
        ('Recuperación IFO×época en run02v2',
         'Runtime recovery: rango GPS → O3a/O3b; pertenencia a CSV H1 → H1 vs. L1.'),
        ('Conflicto m2_anfis.py vs. paquete m2_anfis/',
         'Carga dinámica con importlib.util.spec_from_file_location().'),
        ('Desequilibrio de clase Other (F1=0,36)',
         'Class-weights inversamente proporcionales a frecuencia en CrossEntropy.'),
    ]:
        bullet(doc, lbl, txt)

    h(doc, '5.4 Posibilidades de escalado', 2)

    h(doc, '5.4.1 Eje 1 — Escalado de datos', 3)
    p(doc,
      'Ampliar a O4 (en curso desde mayo 2023), incorporar el detector Virgo (V1) '
      'y actualizar los CSV de Gravity Spy O4 elevaría el dataset etiquetado de '
      '~1700 a >10 000 muestras, reduciendo el desequilibrio de clases y '
      'permitiendo radios ra menores (más reglas ANFIS) sin riesgo de sobreajuste.')

    h(doc, '5.4.2 Eje 2 — Escalado arquitectónico', 3)
    p(doc,
      'El diseño original contempla M3 (regresión causal con canales auxiliares '
      'del interferómetro para identificar la fuente física del glitch) y M4 '
      '(sustracción controlada del ruido). La adición de un lazo RL sobre M2 '
      'permitiría adaptar el clasificador online ante derivas lentas del ruido '
      'durante el run, sin reentrenamiento completo.')

    h(doc, '5.4.3 Eje 3 — Escalado de despliegue', 3)
    p(doc,
      'Adaptar M0 a modo streaming (Kafka/Redis Streams) y desplegar M1+M2 '
      'en GPU (latencia estimada <2 ms) permitiría clasificación en tiempo '
      'cuasi-real. Un dashboard Grafana+InfluxDB con tasa de glitches por clase '
      'y activación de reglas ANFIS podría integrarse en el sistema de alertas '
      'de LIGO-Virgo-KAGRA.')

    h(doc, '5.4.4 Eje 4 — Escalado interpretativo', 3)
    p(doc,
      'La exportación de reglas ANFIS a tablas auditables (unidades físicas '
      'originales) y su visualización interactiva (Streamlit) facilitaría '
      'la validación por físicos sin conocimiento de PyTorch. '
      'La reincorporación de UMAP sobre el embedding latente z permitiría '
      'explorar morfologías no previstas y descubrir nuevas clases '
      'emergentes en futuros runs de observación.')

# ─── Referencias ──────────────────────────────────────────────────────────────

def refs(doc):
    h(doc, 'Referencias')
    for ref in [
        '[1] B. P. Abbott et al. (LIGO Sci. & Virgo Collab.), "Observation of Gravitational Waves from a Binary Black Hole Merger," Phys. Rev. Lett., vol. 116, p. 061102, 2016. doi: 10.1103/PhysRevLett.116.061102.',
        '[2] J. Aasi et al. (LIGO Sci. Collab.), "Advanced LIGO," Class. Quantum Grav., vol. 32, p. 074001, 2015. doi: 10.1088/0264-9381/32/7/074001.',
        '[3] M. Zevin et al., "Gravity Spy: integrating advanced LIGO detector characterization, machine learning, and citizen science," Class. Quantum Grav., vol. 34, p. 064003, 2017. doi: 10.1088/1361-6382/aa5cea.',
        '[4] M. Vallisneri et al., "The LIGO Open Science Center," J. Phys.: Conf. Ser., vol. 610, p. 012021, 2015. doi: 10.1088/1742-6596/610/1/012021.',
        '[5] S. Chatterji et al., "Multiresolution techniques for the detection of gravitational-wave bursts," Class. Quantum Grav., vol. 21, pp. S1809–S1818, 2004. doi: 10.1088/0264-9381/21/20/024.',
        '[6] I. Goodfellow, Y. Bengio, and A. Courville, Deep Learning. MIT Press, 2016.',
        '[7] J.-S. R. Jang, "ANFIS: Adaptive-network-based fuzzy inference system," IEEE Trans. Syst. Man Cybern., vol. 23, pp. 665–685, 1993. doi: 10.1109/21.256541.',
        '[8] S. L. Chiu, "Fuzzy model identification based on cluster estimation," J. Intell. Fuzzy Syst., vol. 2, pp. 267–278, 1994. doi: 10.3233/IFS-1994-2306.',
        '[9] D. P. Kingma and J. Ba, "Adam: A Method for Stochastic Optimization," ICLR 2015. arXiv: 1412.6980.',
        '[10] D. M. Macleod et al., "GWpy: A Python package for gravitational-wave astrophysics," SoftwareX, vol. 13, p. 100657, 2021. doi: 10.1016/j.softx.2021.100657.',
        '[11] D. P. Kingma and M. Welling, "Auto-Encoding Variational Bayes," ICLR 2014. arXiv: 1312.6114.',
        '[12] G. Vajente et al., "Machine-learning nonstationary noise out of gravitational-wave detectors," Phys. Rev. D, vol. 101, p. 042003, 2020. doi: 10.1103/PhysRevD.101.042003.',
        '[13] A. Paszke et al., "PyTorch: An Imperative Style, High-Performance Deep Learning Library," NeurIPS, vol. 32, 2019.',
    ]:
        para = doc.add_paragraph()
        r = para.add_run(ref)
        r.font.name = 'Calibri'; r.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(para, 360, 0, 5)

# ─── Anexo A (3 scripts representativos) ──────────────────────────────────────

def annexe(doc):
    doc.add_page_break()
    h(doc, 'Anexo A. Código fuente (scripts principales)')
    p(doc,
      'Se incluyen los tres scripts más representativos de cada módulo del pipeline. '
      'El resto de archivos se encuentran en el repositorio del proyecto '
      '(run03_bulk_generator.py, run03_minority_generator.py, m1_train_v3.py, '
      'm1_eval_run03.py, m2_feature_extractor.py, m2_pipeline_eval.py, '
      'm1m2_retrain_full.py, etc.).',
      aft=10)

    h(doc, 'A.1  Módulo M0 — Generador de datos nominales', 2)
    code_file(doc, 'A.1', ROOT/'run02_v2_generator.py',
        'Genera run02v2: ventanas Q-transform en GPSs aleatorios uniformes '
        'para 4 combos IFO×época. Reanudable por checkpoints.')

    h(doc, 'A.2  Módulo M1 — Arquitectura GlitchAE', 2)
    code_file(doc, 'A.2', ROOT/'m1_anomaly'/'m1_autoencoder.py',
        'Definición de GlitchAE (encoder Conv2d + FC, decoder ConvTranspose2d) '
        'y cómputo del score de anomalía MSE.')

    h(doc, 'A.3  Módulo M2 — ANFIS jerárquico v3', 2)
    code_file(doc, 'A.3', ROOT/'m2_anfis_v3.py',
        'ANFIS v3: Stage1 Loud/Rest + Stage2 Burst/Scatter/Other, '
        'normalización IFO×época P2/P98, aprendizaje híbrido LSE+Adam.')

# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    print('Generando Memoria_PI_LIGO.docx ...')
    gen_arch()
    doc = setup()
    cover(doc)
    cap1(doc)
    cap2(doc)
    cap3(doc)
    cap4(doc)
    cap5(doc)
    refs(doc)
    annexe(doc)
    page_numbers(doc)
    out = ROOT / 'Memoria_PI_LIGO.docx'
    doc.save(str(out))
    print(f'Saved: {out}  ({out.stat().st_size/1e6:.1f} MB)')

if __name__ == '__main__':
    main()
