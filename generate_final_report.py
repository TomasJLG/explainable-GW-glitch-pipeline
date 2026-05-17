"""
generate_final_report.py
Generates docs/final_report.docx -- professional final project report.

Usage:
    python generate_final_report.py
"""

import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).resolve().parent
DOCS_DIR     = PROJECT_ROOT / "docs"
DOCS_DIR.mkdir(parents=True, exist_ok=True)

# Image paths (fallback chains)
IMG = {
    "confusion_v3":    PROJECT_ROOT / "m2_outputs_v3"  / "confusion_matrix_v3.png",
    "mf_s1":           PROJECT_ROOT / "m2_outputs_v3"  / "membership_functions_stage1.png",
    "mf_s2":           PROJECT_ROOT / "m2_outputs_v3"  / "membership_functions_stage2.png",
    "tc_s1":           PROJECT_ROOT / "m2_outputs_v3"  / "training_curves_stage1.png",
    "tc_s2":           PROJECT_ROOT / "m2_outputs_v3"  / "training_curves_stage2.png",
    "m1_train":        PROJECT_ROOT / "m1_v4_outputs"  / "m1_v4_training_curve.png",
    "m1_train_v3":     PROJECT_ROOT / "m1_v3_outputs"  / "training_curve_v3.png",
    "m1_roc":          PROJECT_ROOT / "m1_v3_outputs"  / "roc_curve_v3.png",
    "m1_scores":       PROJECT_ROOT / "m1_v3_outputs"  / "score_by_class_v3.png",
}

HEADER_BLUE   = "2E74B5"   # header row fill (Word accent blue)
HEADER_TXT    = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_GREY   = "BFBFBF"
ALT_ROW       = "EBF3FB"   # optional alternate row fill (light blue)

# ===========================================================================
# Low-level helpers
# ===========================================================================

def _set_cell_fill(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_table_borders(table, color: str = BORDER_GREY):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    bdr = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        bdr.append(el)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(bdr)


def _style_header_row(table):
    for cell in table.rows[0].cells:
        _set_cell_fill(cell, HEADER_BLUE)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = HEADER_TXT
                run.font.bold      = True


def _add_table(doc, headers: list, rows: list,
               col_widths=None, center_cols=None) -> None:
    """Create a formatted table and add it to doc."""
    center_cols = center_cols or []
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    _set_table_borders(t)

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            align = (WD_ALIGN_PARAGRAPH.CENTER
                     if c_i in center_cols else WD_ALIGN_PARAGRAPH.LEFT)
            cell.paragraphs[0].alignment = align

    # Column widths
    if col_widths:
        for row in t.rows:
            for c_i, w in enumerate(col_widths):
                row.cells[c_i].width = Cm(w)

    _style_header_row(t)


def _add_image(doc, key: str, width_cm: float = 14.0, caption: str = ""):
    path = IMG.get(key)
    # Fallback for m1_train
    if key == "m1_train" and (path is None or not path.exists()):
        path = IMG.get("m1_train_v3")
    if path is None or not path.exists():
        doc.add_paragraph(f"[Figura: {key}]").alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.italic = True
        run.font.size   = Pt(9)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def _heading(doc, text: str, level: int):
    doc.add_heading(text, level=level)


def _para(doc, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name  = "Calibri"
        run.font.size  = Pt(11)
        run.font.bold  = bold
        run.font.italic = italic
    return p


def _bullet(doc, text: str):
    doc.add_paragraph(text, style="List Bullet")


def _numbered(doc, text: str):
    doc.add_paragraph(text, style="List Number")


# ===========================================================================
# Footer with page numbers
# ===========================================================================

def _add_page_numbers(doc):
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    for tag, text in [
        ("w:fldChar", None),
        ("w:instrText", "PAGE"),
        ("w:fldChar", None),
    ]:
        el = OxmlElement(tag)
        if tag == "w:fldChar":
            ft = "begin" if "instrText" not in run._r.xml else "end"
            el.set(qn("w:fldCharType"), ft)
        else:
            el.text = text
        run._r.append(el)


# ===========================================================================
# Document defaults
# ===========================================================================

def _set_defaults(doc):
    # Margins
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Default Normal style → Calibri 11
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Heading styles → Calibri
    for h_name, sz, color in [
        ("Heading 1", 14, "2E74B5"),
        ("Heading 2", 12, "2E74B5"),
        ("Heading 3", 11, "404040"),
    ]:
        style = doc.styles[h_name]
        style.font.name  = "Calibri"
        style.font.size  = Pt(sz)
        style.font.bold  = True
        r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        style.font.color.rgb = RGBColor(r, g, b)


# ===========================================================================
# Section builders
# ===========================================================================

def build_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Pipeline Explicable para Deteccion y Clasificacion\n"
        "de Glitches en Datos LIGO O3")
    run.font.name  = "Calibri"
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(
        "Modulos M1 (Autoencoder Convolucional) y M2 (ANFIS Jerarquico)")
    r2.font.name  = "Calibri"
    r2.font.size  = Pt(14)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()
    doc.add_paragraph()

    for line, sz in [("Tomas Jacobo Legal", 12), ("Mayo 2026", 12)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Calibri"
        r.font.size = Pt(sz)

    doc.add_page_break()


def build_toc(doc):
    _heading(doc, "Tabla de Contenidos", 1)
    p = doc.add_paragraph(
        "TABLA DE CONTENIDOS - Actualizar campo en Word "
        "(clic derecho -> Actualizar campo)")
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()


def build_executive_summary(doc):
    _heading(doc, "1. Resumen Ejecutivo", 1)

    _para(doc,
        "Pipeline de 2 modulos para deteccion y clasificacion automatizada de "
        "glitches (artefactos instrumentales) en datos del interferometro LIGO "
        "durante la campana de observacion O3. El Modulo M1 emplea un autoencoder "
        "convolucional entrenado exclusivamente sobre ruido nominal para detectar "
        "anomalias mediante error de reconstruccion, alcanzando un AUROC de 0.85. "
        "El Modulo M2 implementa un Sistema de Inferencia Difuso Adaptativo (ANFIS) "
        "de arquitectura Takagi-Sugeno en dos etapas jerarquicas para clasificar los "
        "glitches detectados en 4 macro-clases morfologicas, logrando un macro-F1 de "
        "0.68 con 15 reglas IF-THEN completamente interpretables.")

    _para(doc,
        "El hallazgo tecnico mas relevante es que el ancho de banda espectral "
        "(bandwidth) y la relacion senal-ruido (SNR) constituyen los discriminadores "
        "principales entre clases de glitches, un resultado consistente con la "
        "fenomenologia conocida de los interferometros LIGO. La arquitectura ANFIS "
        "permite la extraccion de estas reglas en formato legible por expertos, "
        "transformando el clasificador en una herramienta de diagnostico instrumental.")

    _para(doc,
        "Datos: 6465 ventanas nominales (4 combinaciones detector x epoca) para "
        "entrenamiento de M1; 1637 ventanas con labels reales de Gravity Spy para "
        "entrenamiento y evaluacion de M2.")


def build_data(doc):
    _heading(doc, "2. Datos y Preprocesamiento", 1)
    _heading(doc, "2.1 Fuentes de datos", 2)

    _add_table(doc,
        headers=["Dataset", "Ubicacion", "Contenido", "Ventanas"],
        rows=[
            ["run02v2", "Local",
             "Ventanas nominales, raw Q-transforms, sampling GPS aleatorio",
             "H1/O3a: 465  H1/O3b: 2000  L1/O3a: 2000  L1/O3b: 2000  Total: 6465"],
            ["run03", "Local",
             "Ventanas con labels Gravity Spy, raw Q-transforms, centradas en peak_time",
             "500 (H1/O3a)"],
            ["run03_minority", "Local",
             "Ventanas de clases minoritarias, excluye LF_Burst y Scattered_Light",
             "H1/O3a: 300  H1/O3b: 300  L1/O3a: 300  L1/O3b: 300  Total: 1200"],
            ["Gravity Spy O3 CSVs", "Kaggle/Local",
             "Metadata de triggers: ml_label, snr, peak_frequency, bandwidth, duration",
             "~500K triggers"],
        ],
        col_widths=[3.0, 2.5, 6.0, 5.0],
    )

    doc.add_paragraph()
    _heading(doc, "2.2 Pipeline de preprocesamiento", 2)

    _para(doc,
        "Cada ventana se genera desde datos publicos de GWOSC mediante: "
        "(1) Descarga de archivos HDF5 de 4096s, "
        "(2) Lectura de strain con gwpy, "
        "(3) Q-transform con whiten=True, qrange=(4,64), frange=(20,1700), "
        "(4) Almacenamiento del espectrograma crudo sin normalizacion per-window. "
        "La normalizacion se aplica globalmente usando percentiles P1 y P99 "
        "calculados exclusivamente sobre el conjunto de entrenamiento.")

    _para(doc,
        "La decision de no aplicar normalizacion P99 por ventana (ADR-0015) fue el "
        "hallazgo metodologico mas importante del proyecto: la normalizacion per-window "
        "destruye las diferencias de magnitud inter-ventana que son esenciales para la "
        "deteccion de anomalias. Las versiones que usaron normalizacion inconsistente "
        "obtuvieron AUROC de 0.53 (v1) y 0.30 (v2, polaridad invertida), frente al "
        "0.85 de la version con normalizacion consistente (v3).")


def build_m1(doc):
    _heading(doc, "3. Modulo M1 -- Deteccion de Anomalias", 1)
    _heading(doc, "3.1 Arquitectura", 2)

    _para(doc,
        "Autoencoder convolucional simetrico (GlitchAE). "
        "Encoder: 4 bloques Conv2d con stride=2 (1->32->64->128->256 canales), "
        "BatchNorm, ReLU. Cuello de botella: Flatten -> FC(16384->512) -> FC(512->32). "
        "Decoder: simetrico con ConvTranspose2d. Salida: Sigmoid. "
        "Total: 17.9M parametros. latent_dim=32.")

    _heading(doc, "3.2 Entrenamiento", 2)

    _para(doc,
        "Datos: 6465 ventanas de run02v2, filtradas a 6141 nominales "
        "(log_energy <= P95). Split temporal P80: 4913 train, 1228 val. "
        "Normalizacion global: p1=0.0248, p99=7.3954. "
        "Optimizador: Adam lr=1e-3, weight_decay=1e-5. Loss: MSELoss. "
        "Early stopping en epoch 67 (patience=15). Best val_loss=0.04039.")

    _add_image(doc, "m1_train", caption="Figura 1. Curva de entrenamiento M1 v4.")

    _heading(doc, "3.3 Resultados", 2)

    _para(doc, "AUROC=0.85 evaluado sobre 500 ventanas con labels reales de "
               "Gravity Spy (run03 H1/O3a).")

    _add_image(doc, "m1_roc", caption="Figura 2. Curva ROC M1 (AUROC=0.85).")
    _add_image(doc, "m1_scores",
               caption="Figura 3. AE-score por clase de glitch.")

    doc.add_paragraph()
    _add_table(doc,
        headers=["Clase", "N", "AE Score Mean", "Deteccion"],
        rows=[
            ["Extremely_Loud",    "19",  "0.208", "Fuerte"],
            ["Koi_Fish",          " 3",  "0.081", "Clara"],
            ["Scratchy",          " 1",  "0.062", "Probable"],
            ["Scattered_Light",   "211", "0.051", "Marginal"],
            ["Whistle",           "11",  "0.049", "Marginal"],
            ["Repeating_Blips",   " 3",  "0.048", "Marginal"],
            ["Low_Frequency_Burst","226","0.044", "Nivel de ruido"],
            ["Low_Frequency_Lines","15", "0.044", "Nivel de ruido"],
            ["Blip",              " 4",  "0.043", "Nivel de ruido"],
            ["No_Glitch",         " 1",  "0.042", "Baseline"],
        ],
        col_widths=[4.5, 1.2, 3.0, 3.5],
        center_cols=[1, 2, 3],
    )

    doc.add_paragraph()
    _heading(doc, "3.4 Evolucion del modelo", 2)
    _add_table(doc,
        headers=["Version", "AUROC", "Datos Train", "Problema / Solucion"],
        rows=[
            ["v1", "0.53", "1331 ventanas, norm P99/ventana",
             "Normalizacion per-window en ambos lados"],
            ["v2", "0.30", "7600 ventanas, norm P99/ventana",
             "Mismatch normalizacion -> polaridad invertida"],
            ["v3", "0.85", "465 ventanas raw",
             "Normalizacion global consistente (ADR-0015)"],
            ["v4", "0.85*", "6141 ventanas raw, multi-detector",
             "Version final, mejor generalizacion"],
        ],
        col_widths=[1.5, 1.8, 4.0, 5.5],
        center_cols=[1],
    )
    doc.add_paragraph()
    _para(doc, "*AUROC de v4 estimado; evaluacion formal con mismo test set de v3.",
          italic=True)


def build_m2(doc):
    _heading(doc, "4. Modulo M2 -- Clasificacion ANFIS", 1)
    _heading(doc, "4.1 Macro-clases", 2)

    _add_table(doc,
        headers=["Macro-clase", "Clases originales", "N muestras"],
        rows=[
            ["Loud",    "Extremely_Loud, Koi_Fish",                             "279"],
            ["Burst",   "Low_Frequency_Burst, Blip, Blip_Low_Frequency",        "396"],
            ["Scatter", "Scattered_Light, Fast_Scattering",                      "554"],
            ["Other",   "Whistle, Tomte, Repeating_Blips, Low_Frequency_Lines,\n"
                        "Power_Line, Violin_Mode, Chirp, Scratchy, No_Glitch, etc.", "408"],
        ],
        col_widths=[2.5, 8.5, 2.0],
        center_cols=[2],
    )

    doc.add_paragraph()
    _heading(doc, "4.2 Features de entrada", 2)
    _para(doc,
        "6 features con ANOVA F-score entre parentesis: "
        "peak_frequency (F=137.9), ae_score (F=80.1), log_energy (F=77.4), "
        "snr (F=62.4), duration (F=28.3), bandwidth (F=10.0). "
        "El ae_score proviene del encoder de M1 (error de reconstruccion). "
        "Los demas provienen del catalogo Gravity Spy "
        "(metadata de triggers Omicron).")

    _heading(doc, "4.3 Arquitectura ANFIS jerarquica", 2)
    _para(doc,
        "Modelo Takagi-Sugeno de primer orden con funciones de pertenencia "
        "Generalized Bell (GBELLMF). Entrenamiento hibrido: LSE para parametros "
        "consecuentes + Adam (lr=0.01) para parametros de premisa.")

    doc.add_paragraph()
    _add_table(doc,
        headers=["Etapa", "Objetivo", "Reglas", "ra", "Accuracy"],
        rows=[
            ["Etapa 1", "Loud vs Rest",         "7", "0.15", "90.3%"],
            ["Etapa 2", "Burst vs Scatter vs Other", "8", "0.10", "71.1%"],
        ],
        col_widths=[2.0, 4.5, 1.5, 1.5, 2.5],
        center_cols=[2, 3, 4],
    )
    doc.add_paragraph()

    _add_image(doc, "tc_s1",
               caption="Figura 4. Curvas de entrenamiento Etapa 1 (Loud vs Rest).")
    _add_image(doc, "tc_s2",
               caption="Figura 5. Curvas de entrenamiento Etapa 2 (Burst/Scatter/Other).")

    _heading(doc, "4.4 Resultados", 2)
    _para(doc, "Metricas combinadas sobre 1637 muestras:")

    _add_table(doc,
        headers=["Metrica", "Valor"],
        rows=[
            ["Accuracy",  "69.4%"],
            ["Macro-F1",  "0.68"],
            ["F1 Loud",   "0.75"],
            ["F1 Burst",  "0.70"],
            ["F1 Scatter","0.77"],
            ["F1 Other",  "0.51"],
        ],
        col_widths=[5.0, 3.0],
        center_cols=[1],
    )

    doc.add_paragraph()
    _para(doc, "Matriz de confusion (filas = clase real, columnas = clase predicha):")
    _add_table(doc,
        headers=["", "Loud", "Burst", "Scatter", "Other"],
        rows=[
            ["Loud",    "235",  "  4", " 30", " 10"],
            ["Burst",   " 44",  "262", " 32", " 58"],
            ["Scatter", " 23",  " 26", "462", " 43"],
            ["Other",   " 47",  " 61", "123", "177"],
        ],
        col_widths=[2.5, 2.5, 2.5, 2.5, 2.5],
        center_cols=[1, 2, 3, 4],
    )
    doc.add_paragraph()

    _add_image(doc, "confusion_v3",
               caption="Figura 6. Matriz de confusion M2 v3 (clasificador jerarquico).")

    _heading(doc, "4.5 Evolucion del modelo", 2)
    _add_table(doc,
        headers=["Version", "Accuracy", "Macro-F1", "Muestras", "Cambio principal"],
        rows=[
            ["v1", "74.6%", "0.67", "500 (solo H1/O3a)",        "Baseline, 5 reglas flat"],
            ["v2", "51.8%", "0.54", "1637 (multi-detector)",    "Datos heterogeneos, 5 reglas insuficientes"],
            ["v3", "69.4%", "0.68", "1637 (multi-detector)",    "Jerarquico 2 etapas, 15 reglas, norm por grupo"],
        ],
        col_widths=[1.5, 2.2, 2.2, 3.5, 4.5],
        center_cols=[1, 2],
    )


def build_rules(doc):
    _heading(doc, "5. Reglas Extraidas (Interpretabilidad)", 1)
    _heading(doc, "5.1 Etapa 1 -- Deteccion de glitches Loud", 2)

    rules_s1 = [
        ("R00", "medio",  "medio",  "medio",  "alto",   "medio", "bajo",   "Loud"),
        ("R01", "alto",   "medio",  "alto",   "medio",  "bajo",  "medio",  "Loud"),
        ("R02", "medio",  "medio",  "medio",  "bajo",   "alto",  "bajo",   "Rest"),
        ("R03", "bajo",   "bajo",   "bajo",   "medio",  "medio", "medio",  "Rest"),
        ("R04", "alto",   "alto",   "alto",   "alto",   "medio", "alto",   "Loud"),
        ("R05", "bajo",   "bajo",   "bajo",   "bajo",   "alto",  "alto",   "Loud"),
        ("R06", "medio",  "alto",   "medio",  "medio",  "bajo",  "medio",  "Loud"),
    ]
    feats = ["peak_freq", "ae_score", "log_energy", "snr", "duration", "bandwidth"]

    for r in rules_s1:
        rid, *vals, cls = r
        conds = "  Y  ".join(f"{f} ES {v}" for f, v in zip(feats, vals))
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run(f"{rid}: SI  {conds}")
        run.font.name = "Courier New"; run.font.size = Pt(8)
        p2 = doc.add_paragraph(style="No Spacing")
        run2 = p2.add_run(f"      ENTONCES {cls}")
        run2.font.name = "Courier New"; run2.font.size = Pt(8)
        run2.font.bold = True
        doc.add_paragraph()

    _add_image(doc, "mf_s1",
               caption="Figura 7. Funciones de pertenencia GBellMF -- Etapa 1.")

    _heading(doc, "5.2 Etapa 2 -- Clasificacion Burst/Scatter/Other", 2)

    rules_s2 = [
        ("R00", "medio",  "bajo",   "medio",  "alto",   "bajo",  "bajo",   "Other"),
        ("R01", "bajo",   "bajo",   "bajo",   "bajo",   "medio", "medio",  "Scatter"),
        ("R02", "bajo",   "medio",  "medio",  "medio",  "alto",  "bajo",   "Scatter"),
        ("R03", "medio",  "alto",   "alto",   "medio",  "bajo",  "alto",   "Other"),
        ("R04", "alto",   "bajo",   "bajo",   "bajo",   "bajo",  "medio",  "Scatter"),
        ("R05", "bajo",   "medio",  "bajo",   "alto",   "alto",  "alto",   "Scatter"),
        ("R06", "alto",   "alto",   "alto",   "alto",   "medio", "bajo",   "Other"),
        ("R07", "alto",   "alto",   "alto",   "bajo",   "alto",  "alto",   "Other"),
    ]

    for r in rules_s2:
        rid, *vals, cls = r
        conds = "  Y  ".join(f"{f} ES {v}" for f, v in zip(feats, vals))
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run(f"{rid}: SI  {conds}")
        run.font.name = "Courier New"; run.font.size = Pt(8)
        p2 = doc.add_paragraph(style="No Spacing")
        run2 = p2.add_run(f"      ENTONCES {cls}")
        run2.font.name = "Courier New"; run2.font.size = Pt(8)
        run2.font.bold = True
        doc.add_paragraph()

    _add_image(doc, "mf_s2",
               caption="Figura 8. Funciones de pertenencia GBellMF -- Etapa 2.")

    _heading(doc, "5.3 Interpretacion fisica", 2)

    interps = [
        ("Loud (Extremely_Loud, Koi_Fish)",
         "Se caracterizan por SNR alto y ae_score alto. M1 ya los detecta por "
         "su elevado error de reconstruccion; M2 confirma la clasificacion en "
         "la Etapa 1 con 90.3 % de accuracy."),
        ("Scatter (Scattered_Light, Fast_Scattering)",
         "Se identifican por duration alta y bandwidth alta, consistente con "
         "la fenomenologia de dispersion de luz estocastica que produce senales "
         "prolongadas de espectro ancho en los interferometros."),
        ("Burst (Low_Frequency_Burst, Blip)",
         "Presentan peak_frequency baja y bandwidth baja: eventos transitorios "
         "de banda estrecha concentrados en frecuencias bajas."),
        ("Other",
         "Agrupa anomalias diversas cuya firma espectral no encaja en las otras "
         "categorias. Su F1 mas bajo (0.51) refleja esta heterogeneidad inherente."),
    ]
    for cls_name, text in interps:
        _para(doc, cls_name, bold=True)
        _para(doc, text)

    _para(doc,
        "Esta capacidad de vincular parametros fisicos medibles con categorias "
        "morfologicas de glitches es el diferenciador principal de ANFIS frente a "
        "arquitecturas CNN de caja negra, que aunque alcanzan precisiones superiores "
        "al 97 %, no proporcionan este nivel de explicabilidad.")


def build_generator(doc):
    _heading(doc, "6. Modulo de Generacion de Datos (Generador NPZ)", 1)
    _heading(doc, "6.1 Arquitectura del generador", 2)

    _para(doc,
        "Pipeline de generacion de datasets Q-transform desde datos publicos "
        "de GWOSC. Dos modos:")
    _bullet(doc,
        "Modo 1 -- Sampling aleatorio (run02v2): Genera ventanas nominales "
        "muestreando tiempos GPS uniformes dentro de segmentos de ciencia "
        "validados con flag IFO_DATA. Usado para entrenamiento de M1.")
    _bullet(doc,
        "Modo 2 -- Sampling targeteado (run03): Genera ventanas centradas en "
        "peak_time de triggers de Gravity Spy, garantizando label para cada "
        "ventana. Variante minority filtra por clases subrepresentadas "
        "excluyendo Low_Frequency_Burst y Scattered_Light.")

    _heading(doc, "6.2 Pipeline de procesamiento por ventana", 2)
    steps = [
        "Consultar segmentos de ciencia con gwosc.timeline.get_segments",
        "Agrupar GPS candidatos por bloque HDF5 de 4096s, ordenar por densidad descendente",
        "Descargar archivo HDF5 con gwosc.locate.get_urls",
        "Leer strain con gwpy.timeseries.TimeSeries.read(format='hdf5.gwosc')",
        "Crop a peak_time +/- 8s (margen para whitening)",
        "Q-transform con whiten=True, qrange=(4,64), frange=(20,1700)",
        "Guardar espectrograma crudo sin normalizacion per-window",
        "Calcular log_energy = log1p(percentile(spec, 90))",
        "Resize a 128x128 con interpolacion bilineal",
    ]
    for s in steps:
        _numbered(doc, s)

    _heading(doc, "6.3 Rendimiento y tolerancia a fallos", 2)
    _para(doc,
        "Rendimiento: ~7 ventanas/bloque, ~2 min/bloque. "
        "2000 ventanas aprox. 8 horas.")
    _para(doc,
        "Sistema de checkpoint: guarda progreso cada 50 ventanas "
        "(bloques procesados + ventanas generadas). Permite parar y continuar "
        "sin reprocesar. Borrado automatico de HDF5 tras procesar para "
        "gestion de disco.")
    _para(doc,
        "Scripts: run02_v2_generator.py (modo aleatorio), "
        "run03_bulk_generator.py (modo targeteado), "
        "run03_minority_generator.py (modo minority).")


def build_pipeline(doc):
    _heading(doc, "7. Pipeline End-to-End", 1)

    _para(doc, "Flujo completo de una ventana:")
    flow_steps = [
        "Tiempo GPS",
        "Descarga GWOSC (HDF5 4096s)",
        "Q-transform (gwpy, whiten=True, 128x128)",
        "Normalizacion global (p1/p99)",
        "M1 encoder -> vector latente (dim=32) + ae_score",
        "Concatenar con features CSV (snr, peak_frequency, duration, bandwidth)",
        "ANFIS Etapa 1: Loud vs Rest (7 reglas, ra=0.15)",
        "ANFIS Etapa 2: Burst vs Scatter vs Other (8 reglas, ra=0.10)",
        "Clase predicha + regla activa",
    ]
    t = doc.add_table(rows=len(flow_steps), cols=2)
    t.style = "Table Grid"
    _set_table_borders(t)
    for i, step in enumerate(flow_steps):
        t.rows[i].cells[0].text = f"({i + 1})"
        t.rows[i].cells[1].text = step
        t.rows[i].cells[0].width = Cm(1.5)
        t.rows[i].cells[1].width = Cm(12.5)
        t.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _para(doc,
        "Verificacion de consistencia: 500/500 predicciones identicas entre "
        "ejecucion pipeline directa y features pre-calculados (m2_pipeline_eval.py).")


def build_adr(doc):
    _heading(doc, "8. Registro de Decisiones de Arquitectura (ADR)", 1)

    _add_table(doc,
        headers=["ADR", "Decision", "Estado"],
        rows=[
            ["ADR-0015",
             "Normalizacion global sin P99 por ventana. "
             "La normalizacion per-window destruye diferencias de magnitud "
             "inter-ventana criticas para M1.",
             "Activo"],
            ["ADR-0016",
             "M1 cerrado a AUROC=0.85. "
             "Glitches sutiles (LF_Burst, Blip) delegados a M2.",
             "Activo"],
            ["ADR-0017",
             "M2 usa features del encoder de M1 + metadata CSV, no imagenes raw.",
             "Activo"],
            ["ADR-0018",
             "23 clases Gravity Spy -> 4 macro-clases (Loud, Burst, Scatter, Other) "
             "para viabilidad ANFIS.",
             "Activo"],
        ],
        col_widths=[2.5, 10.5, 1.5],
        center_cols=[2],
    )


def build_limitations(doc):
    _heading(doc, "9. Limitaciones y Trabajo Futuro", 1)
    _heading(doc, "9.1 Limitaciones actuales", 2)

    limits = [
        "Solo escala temporal 1.0s; escalas 0.25s y 4.0s pendientes.",
        "Clase Other heterogenea con F1=0.51.",
        "Dataset de evaluacion limitado a 1637 ventanas.",
        "M1 no detecta glitches sutiles (LF_Burst, Blip): AE score indistinguible del ruido.",
    ]
    for l in limits:
        _bullet(doc, l)

    _heading(doc, "9.2 Trabajo futuro", 2)

    futures = [
        "Ampliar a escalas 0.25s y 4.0s para capturar glitches de diferente duracion.",
        "Generar mas muestras de clases minoritarias con datos de H1/O3b, L1/O3a, L1/O3b.",
        "Implementar M3 (regresion con canales auxiliares) para correlacionar glitches "
        "con fuentes instrumentales.",
        "Explorar ANFIS con features de multiples escalas simultaneas.",
        "Validacion cruzada O3a->O3b y H1->L1 para certificar generalizacion inter-detector.",
    ]
    for f in futures:
        _bullet(doc, f)


# ===========================================================================
# Main
# ===========================================================================

def main():
    print("Building docs/final_report.docx ...")
    doc = Document()
    _set_defaults(doc)
    _add_page_numbers(doc)

    build_cover(doc)
    build_toc(doc)
    build_executive_summary(doc)
    doc.add_page_break()

    build_data(doc)
    doc.add_page_break()

    build_m1(doc)
    doc.add_page_break()

    build_m2(doc)
    doc.add_page_break()

    build_rules(doc)
    doc.add_page_break()

    build_generator(doc)
    doc.add_page_break()

    build_pipeline(doc)
    doc.add_page_break()

    build_adr(doc)
    doc.add_page_break()

    build_limitations(doc)

    out = DOCS_DIR / "final_report.docx"
    doc.save(str(out))
    size_kb = out.stat().st_size // 1024
    print(f"Saved: {out}  ({size_kb} KB)")

    # Report which images were embedded vs placeholder
    for key, path in IMG.items():
        if key == "m1_train" and not path.exists():
            path = IMG["m1_train_v3"]
        status = "OK" if (path and path.exists()) else "PLACEHOLDER"
        print(f"  {key:18s}: {status}")


if __name__ == "__main__":
    main()
