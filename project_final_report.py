"""
project_final_report.py
Generates docs/final_report.docx -- complete M1+M2 pipeline final report.

Usage:
    python project_final_report.py
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).resolve().parent
DOCS_DIR     = PROJECT_ROOT / "docs"
M1_OUT       = PROJECT_ROOT / "m1_v3_outputs"
M2_OUT       = PROJECT_ROOT / "m2_outputs"
M2_PIPE      = M2_OUT / "pipeline_eval"

DOCS_DIR.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _style_header_row(table, hex_color: str = "1F3864"):
    row = table.rows[0]
    for cell in row.cells:
        _set_cell_bg(cell, hex_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold      = True


def _add_table(doc, headers: list, rows: list, col_widths=None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
    if col_widths:
        for r in t.rows:
            for c_idx, w in enumerate(col_widths):
                r.cells[c_idx].width = Inches(w)
    _style_header_row(t)


def _add_image(doc, path: Path, width_in: float = 5.5, caption: str = ""):
    if not path.exists():
        doc.add_paragraph(f"[Imagen no encontrada: {path.name}]")
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.italic = True
        p.runs[0].font.size   = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _h1(doc, text):
    doc.add_heading(text, level=1)


def _h2(doc, text):
    doc.add_heading(text, level=2)


def _h3(doc, text):
    doc.add_heading(text, level=3)


def _para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.font.bold = True
    return p


def _bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def build_title(doc):
    doc.add_heading(
        "Pipeline Explicable para Deteccion y Clasificacion de "
        "Glitches en Datos LIGO O3",
        level=0,
    )
    p = doc.add_paragraph(
        "Proyecto M1-M2 | Detector de Anomalias\n"
        "Informe Final"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def build_executive_summary(doc):
    _h1(doc, "1. Resumen Ejecutivo")
    _para(doc,
        "Este proyecto implementa un pipeline de dos modulos para la deteccion "
        "y clasificacion explicable de artefactos instrumentales (glitches) en "
        "los datos del detector LIGO H1 durante la corrida O3a de ondas "
        "gravitacionales."
    )

    _h2(doc, "Pipeline General")
    _bullet(doc, "M1 -- Deteccion no supervisada: Autoencoder convolucional "
                 "(GlitchAE) entrenado sobre ventanas nominales. "
                 "AUROC = 0.85.")
    _bullet(doc, "M2 -- Clasificacion interpretable: Clasificador ANFIS "
                 "Takagi-Sugeno de 1er orden con 5 reglas difusas. "
                 "Accuracy = 74.6 %, Macro-F1 = 0.67.")
    _bullet(doc, "Validacion end-to-end: 500/500 muestras con consistencia "
                 "perfecta entre inferencia de pipeline y features precalculadas.")

    _h2(doc, "Hallazgos Clave")
    _bullet(doc,
        "Bandwidth es el discriminador principal entre clases: ancho de banda "
        "estrecho -> Scatter; ancho de banda medio/alto -> Burst/Other."
    )
    _bullet(doc,
        "Extremely_Loud y Koi_Fish son altamente detectables por M1 "
        "(mediana AE-score > 0.08); el resto queda cerca del umbral nominal."
    )
    _bullet(doc,
        "El cuello de botella del pipeline es el tamano del dataset: 500 "
        "muestras de evaluacion con ~90 % en solo dos clases (Burst + Scatter)."
    )

    _h2(doc, "Resumen de Metricas")
    _add_table(
        doc,
        headers=["Modulo", "Metrica", "Valor"],
        rows=[
            ["M1 GlitchAE",   "AUROC",           "0.85"],
            ["M1 GlitchAE",   "Best val loss",    "0.0420"],
            ["M1 GlitchAE",   "Parametros",       "17.9 M"],
            ["M2 ANFIS",      "Accuracy (500)",   "74.6 %"],
            ["M2 ANFIS",      "Macro-F1 (500)",   "0.67"],
            ["M2 ANFIS",      "Accuracy (test)",  "68.6 %"],
            ["M2 ANFIS",      "Macro-F1 (test)",  "0.60"],
            ["M2 ANFIS",      "Reglas difusas",   "5"],
            ["M2 ANFIS",      "Parametros",       "230"],
        ],
        col_widths=[1.8, 2.2, 1.5],
    )
    doc.add_paragraph()


def build_m1(doc):
    _h1(doc, "2. Modulo M1 -- Deteccion de Anomalias")

    _h2(doc, "2.1 Arquitectura")
    _add_table(
        doc,
        headers=["Componente", "Detalle"],
        rows=[
            ["Modelo",           "GlitchAE (autoencoder convolucional)"],
            ["latent_dim",       "32"],
            ["Parametros",       "17.9 M"],
            ["Input",            "Q-transform crudo (H1, O3a, 1 s)"],
            ["Normalizacion",    "Global P1/P99 (ADR-0015)"],
            ["P1",               "0.0250"],
            ["P99",              "7.3544"],
            ["Loss",             "MSE reconstruccion"],
            ["Epochs",           "80 (convergencia)"],
            ["Best val loss",    "0.0420"],
        ],
        col_widths=[2.0, 4.0],
    )
    doc.add_paragraph()

    _h2(doc, "2.2 Datos de Entrenamiento")
    _bullet(doc, "465 ventanas nominales (run02v2, H1/O3a, 1 s)")
    _bullet(doc, "80/20 train/val split")
    _bullet(doc, "Sin glitches en entrenamiento (aprendizaje no supervisado)")

    _h2(doc, "2.3 Resultados: AUROC = 0.85")
    _add_image(doc, M1_OUT / "roc_curve_v3.png", 5.0,
               "Figura 1. Curva ROC del modulo M1 (AUROC = 0.85).")
    doc.add_paragraph()

    _add_image(doc, M1_OUT / "score_distribution_v3.png", 5.0,
               "Figura 2. Distribucion de AE-scores (nominal vs glitch).")
    doc.add_paragraph()

    _add_image(doc, M1_OUT / "score_by_class_v3.png", 5.5,
               "Figura 3. AE-score por clase de glitch.")
    doc.add_paragraph()

    _h2(doc, "2.4 Detectabilidad por Clase")
    _para(doc,
        "La siguiente tabla muestra el AE-score mediano y medio para cada "
        "clase en las 500 muestras de evaluacion. Un score mas alto indica "
        "mayor divergencia de la reconstruccion nominal y, por tanto, mayor "
        "detectabilidad."
    )
    doc.add_paragraph()
    _add_table(
        doc,
        headers=["Clase (original)", "Macro-clase", "n", "Mediana AE", "Media AE", "Detectabilidad"],
        rows=[
            ["Extremely_Loud",      "Loud",    "19", "0.113", "0.208", "Alta"],
            ["Koi_Fish",            "Loud",    " 3", "0.084", "0.081", "Media"],
            ["Scattered_Light",     "Scatter", "211","0.050", "0.051", "Baja"],
            ["Fast_Scattering",     "Scatter", " 2", "0.050", "0.050", "Baja"],
            ["Low_Frequency_Burst", "Burst",   "226","0.044", "0.044", "Baja"],
            ["Blip",                "Burst",   " 4", "0.043", "0.043", "Baja"],
            ["Blip_Low_Frequency",  "Burst",   " 3", "0.042", "0.045", "Baja"],
            ["Low_Frequency_Lines", "Line",    "15", "0.044", "0.044", "Baja"],
            ["Whistle",             "Other",   "11", "0.047", "0.049", "Baja"],
            ["Repeating_Blips",     "Other",   " 3", "0.048", "0.048", "Baja"],
            ["Scratchy",            "Other",   " 1", "0.062", "0.062", "Media"],
            ["No_Glitch",           "Other",   " 1", "0.042", "0.042", "Baja"],
            ["Chirp",               "Other",   " 1", "0.050", "0.050", "Baja"],
        ],
        col_widths=[2.2, 1.2, 0.5, 1.1, 1.0, 1.3],
    )
    doc.add_paragraph()

    _h2(doc, "2.5 Curva de Entrenamiento")
    _add_image(doc, M1_OUT / "training_curve_v3.png", 5.0,
               "Figura 4. Curva de loss de entrenamiento y validacion M1.")
    doc.add_paragraph()

    _h2(doc, "2.6 Leccion Aprendida")
    _para(doc,
        "La consistencia de la normalizacion entre entrenamiento y evaluacion "
        "es critica. Una normalizacion por ventana (P99 local) destruye la "
        "senal de anomalia porque los glitches intensos quedan renormalizados "
        "a valores similares a las ventanas nominales. La solucion (ADR-0015) "
        "fue usar estadisticas globales P1/P99 calculadas sobre el conjunto "
        "de entrenamiento nominal y almacenadas en normalization_v3.json."
    )


def build_m2(doc):
    _h1(doc, "3. Modulo M2 -- Clasificacion ANFIS")

    _h2(doc, "3.1 Arquitectura")
    _add_table(
        doc,
        headers=["Componente", "Detalle"],
        rows=[
            ["Modelo",          "ANFIS Takagi-Sugeno 1er orden"],
            ["MF",              "GBellMF (campana generalizada)"],
            ["Reglas",          "5 (subtractive clustering, ra=0.15)"],
            ["Features (6)",    "peak_frequency, ae_score, log_energy, snr, duration, bandwidth"],
            ["Clases (4)",      "Loud, Burst, Scatter, Other"],
            ["Parametros",      "90 premisa + 140 consecuente = 230 total"],
            ["Entrenamiento",   "Hibrido LSE + Adam"],
            ["Max epochs",      "200 (early stop patience=20)"],
            ["Datos",           "500 muestras run03, 80/20 train/test, seed=42"],
        ],
        col_widths=[2.5, 4.5],
    )
    doc.add_paragraph()

    _h2(doc, "3.2 Macro-clases")
    _para(doc,
        "Las 23 clases originales de Gravity Spy se agrupan en 4 macro-clases "
        "para el clasificador (ADR-0018). Line se fusiona con Other por ser "
        "poco representada (15 muestras total)."
    )
    _add_table(
        doc,
        headers=["Macro-clase", "Clases originales", "n (500 total)"],
        rows=[
            ["Loud",    "Extremely_Loud, Koi_Fish",                              "22"],
            ["Burst",   "Low_Frequency_Burst, Blip, Blip_Low_Frequency",         "233"],
            ["Scatter", "Scattered_Light, Fast_Scattering",                       "213"],
            ["Other",   "Whistle, Repeating_Blips, Chirp, Scratchy, No_Glitch, "
                        "Low_Frequency_Lines (+Line)",                            "32"],
        ],
        col_widths=[1.4, 3.8, 1.1],
    )
    doc.add_paragraph()

    _h2(doc, "3.3 Resultados Globales (500 muestras)")
    _add_table(
        doc,
        headers=["Clase", "Precision", "Recall", "F1", "Soporte"],
        rows=[
            ["Loud",    "0.889", "0.727", "0.800", "22"],
            ["Burst",   "0.702", "0.888", "0.784", "233"],
            ["Scatter", "0.794", "0.671", "0.728", "213"],
            ["Other",   "1.000", "0.219", "0.359", "32"],
            ["MACRO",   "--",    "--",    "0.668", "500"],
        ],
        col_widths=[1.5, 1.2, 1.2, 1.2, 1.2],
    )
    doc.add_paragraph()

    _h2(doc, "3.4 Resultados Test Split (102 muestras, 80/20 seed=42)")
    _add_table(
        doc,
        headers=["Clase", "Precision", "Recall", "F1", "Soporte"],
        rows=[
            ["Loud",    "1.000", "0.600", "0.750", "5"],
            ["Burst",   "0.667", "0.809", "0.731", "47"],
            ["Scatter", "0.683", "0.651", "0.667", "43"],
            ["Other",   "1.000", "0.143", "0.250", "7"],
            ["MACRO",   "--",    "--",    "0.599", "102"],
        ],
        col_widths=[1.5, 1.2, 1.2, 1.2, 1.2],
    )
    doc.add_paragraph()

    _h2(doc, "3.5 Matriz de Confusion (500 muestras)")
    _para(doc, "Filas = clase verdadera, columnas = clase predicha.")
    _add_table(
        doc,
        headers=["", "Loud (pred)", "Burst (pred)", "Scatter (pred)", "Other (pred)"],
        rows=[
            ["Loud (true)",    "16",  " 0",  " 6",  "0"],
            ["Burst (true)",   " 0",  "207", "26",  "0"],
            ["Scatter (true)", " 2",  "68",  "143", "0"],
            ["Other (true)",   " 0",  "20",  " 5",  "7"],
        ],
        col_widths=[1.5, 1.4, 1.4, 1.5, 1.5],
    )
    doc.add_paragraph()

    _add_image(doc, M2_PIPE / "pipeline_confusion_matrix.png", 5.0,
               "Figura 5. Matriz de confusion del pipeline M1+M2 (500 muestras).")
    doc.add_paragraph()

    _h2(doc, "3.6 Importancia de Features")
    _add_image(doc, M2_PIPE / "feature_importance.png", 5.0,
               "Figura 6. Importancia de features por ANOVA F-score.")
    doc.add_paragraph()

    _para(doc,
        "El ranking de importancia (ANOVA F-score sobre 500 muestras) es: "
        "peak_frequency (F=137.9) > ae_score (F=80.1) > log_energy (F=77.4) "
        "> snr (F=62.4) > duration (F=28.3) > bandwidth (F=10.0). "
        "Sin embargo, el analisis de reglas muestra que bandwidth es el "
        "discriminador efectivo en tiempo de inferencia."
    )

    _h2(doc, "3.7 Curva de Entrenamiento")
    _add_image(doc, M2_OUT / "training_curve_m2.png", 5.0,
               "Figura 7. Curva de entrenamiento M2 ANFIS (hibrido LSE+Adam).")
    doc.add_paragraph()

    _h2(doc, "3.8 Funciones de Pertenencia")
    _add_image(doc, M2_OUT / "membership_functions_m2.png", 5.5,
               "Figura 8. Funciones de pertenencia GBELLMF por feature y regla.")
    doc.add_paragraph()


def build_rules(doc):
    _h1(doc, "4. Reglas Extraidas -- Interpretabilidad")

    _para(doc,
        "Las 5 reglas difusas del sistema ANFIS se extraen directamente de "
        "los parametros entrenados. Para cada regla se reportan: los centros "
        "de las MF (c), anchuras (a) y exponentes (b) en escala normalizada "
        "[0,1], la clase dominante del consecuente, y los pesos lineales por "
        "clase."
    )
    doc.add_paragraph()

    # Rules data from pipeline eval output
    rules = [
        {
            "id": "R003", "rank": 1, "avg_strength": 0.2225, "output": "Scatter",
            "key_feature": "bandwidth IS low",
            "mf_params": [
                ("peak_frequency", "low",    "c=-0.028", "a=0.293", "b=2.022"),
                ("ae_score",       "low",    "c=0.021",  "a=0.298", "b=2.008"),
                ("log_energy",     "low",    "c=0.006",  "a=0.298", "b=2.009"),
                ("snr",            "low",    "c=-0.008", "a=0.297", "b=2.009"),
                ("duration",       "low",    "c=0.088",  "a=0.297", "b=2.010"),
                ("bandwidth",      "low",    "c=0.280",  "a=0.296", "b=2.011"),
            ],
            "weights": "Loud=+0.123, Burst=-2.045, Scatter=+4.391, Other=-1.469",
            "interpretation":
                "Glitches con bandwidth estrecho y baja energia. "
                "Correspondencia fisica: luz dispersada (Scattered Light) "
                "genera arcos de frecuencia estrecha en el espectrograma.",
        },
        {
            "id": "R002", "rank": 2, "avg_strength": 0.2141, "output": "Burst",
            "key_feature": "bandwidth IS medium",
            "mf_params": [
                ("peak_frequency", "low",    "c=0.035",  "a=0.308", "b=1.974"),
                ("ae_score",       "low",    "c=0.025",  "a=0.299", "b=2.004"),
                ("log_energy",     "low",    "c=0.001",  "a=0.299", "b=1.997"),
                ("snr",            "low",    "c=-0.002", "a=0.307", "b=1.976"),
                ("duration",       "low",    "c=0.105",  "a=0.307", "b=1.977"),
                ("bandwidth",      "medium", "c=0.622",  "a=0.307", "b=1.977"),
            ],
            "weights": "Loud=-2.231, Burst=+7.665, Scatter=-7.161, Other=+2.727",
            "interpretation":
                "Glitches con ancho de banda medio. Correspondencia fisica: "
                "Low Frequency Burst ocupa una banda espectral moderada "
                "(tipicamente 10-100 Hz de ancho).",
        },
        {
            "id": "R000", "rank": 3, "avg_strength": 0.2132, "output": "Scatter",
            "key_feature": "bandwidth IS low",
            "mf_params": [
                ("peak_frequency", "low",    "c=-0.021", "a=0.294", "b=2.012"),
                ("ae_score",       "low",    "c=0.032",  "a=0.299", "b=2.002"),
                ("log_energy",     "low",    "c=0.010",  "a=0.299", "b=2.001"),
                ("snr",            "low",    "c=-0.000", "a=0.300", "b=2.000"),
                ("duration",       "low",    "c=0.077",  "a=0.300", "b=2.000"),
                ("bandwidth",      "low",    "c=0.002",  "a=0.299", "b=2.004"),
            ],
            "weights": "Loud=-0.035, Burst=-1.165, Scatter=+2.715, Other=-0.515",
            "interpretation":
                "Regla base de Scatter: bandwidth muy estrecho (centro c=0.002 "
                "en escala normalizada). Captura el core de la clase Scattered_Light.",
        },
        {
            "id": "R001", "rank": 4, "avg_strength": 0.1954, "output": "Burst",
            "key_feature": "bandwidth IS high",
            "mf_params": [
                ("peak_frequency", "low",    "c=-0.026", "a=0.293", "b=2.026"),
                ("ae_score",       "low",    "c=0.029",  "a=0.300", "b=1.999"),
                ("log_energy",     "low",    "c=0.012",  "a=0.300", "b=1.999"),
                ("snr",            "low",    "c=0.005",  "a=0.300", "b=2.003"),
                ("duration",       "low",    "c=0.117",  "a=0.299", "b=2.003"),
                ("bandwidth",      "high",   "c=0.771",  "a=0.299", "b=1.998"),
            ],
            "weights": "Loud=-0.189, Burst=+1.925, Scatter=+1.772, Other=-2.508",
            "interpretation":
                "Burst de ancho de banda alto. Competencia entre Burst y "
                "Scatter en los pesos consecuentes, lo que explica la "
                "confusion residual entre ambas clases.",
        },
        {
            "id": "R004", "rank": 5, "avg_strength": 0.1548, "output": "Other",
            "key_feature": "bandwidth IS high",
            "mf_params": [
                ("peak_frequency", "low",    "c=0.027",  "a=0.307", "b=1.975"),
                ("ae_score",       "low",    "c=0.015",  "a=0.300", "b=2.000"),
                ("log_energy",     "low",    "c=0.005",  "a=0.300", "b=2.000"),
                ("snr",            "low",    "c=-0.001", "a=0.300", "b=1.997"),
                ("duration",       "low",    "c=0.160",  "a=0.301", "b=1.997"),
                ("bandwidth",      "high",   "c=0.943",  "a=0.297", "b=2.003"),
            ],
            "weights": "Loud=+0.229, Burst=-0.742, Scatter=-1.565, Other=+3.077",
            "interpretation":
                "Glitches con bandwidth muy alto y duracion ligeramente mayor. "
                "Captura whistles y glitches broadband que no encajan en "
                "Burst ni Scatter.",
        },
    ]

    for rule in rules:
        _h2(doc, f"Regla {rule['id']} -- THEN class = {rule['output']}  "
                 f"(fuerza media = {rule['avg_strength']:.4f})")

        # IF-THEN block
        cond_lines = ["  IF  " + rule["mf_params"][0][0].ljust(17) +
                      f"  IS  {rule['mf_params'][0][1]:6s}  "
                      f"[{rule['mf_params'][0][2]}, {rule['mf_params'][0][3]}, {rule['mf_params'][0][4]}]"]
        for feat, label, c, a, b in rule["mf_params"][1:]:
            cond_lines.append(
                f"  AND  {feat.ljust(17)}  IS  {label:6s}  [{c}, {a}, {b}]"
            )
        cond_lines.append(f"  THEN class = {rule['output']}")
        cond_lines.append(f"       weights: {rule['weights']}")
        block = doc.add_paragraph("\n".join(cond_lines))
        block.style = "No Spacing"
        for run in block.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)

        _para(doc, "Interpretacion: " + rule["interpretation"])
        doc.add_paragraph()

    _h2(doc, "4.1 Bandwidth como Discriminador Principal")
    _para(doc,
        "El analisis de las 5 reglas muestra que bandwidth es el unico "
        "feature que varia sistematicamente entre reglas: las dos reglas "
        "Scatter tienen bandwidth low (c=0.002 y c=0.280), las reglas Burst "
        "tienen bandwidth medium o high (c=0.622, c=0.771), y la regla Other "
        "tiene el bandwidth mas alto (c=0.943). Este patron es fisicamente "
        "consistente: los glitches de luz dispersada (Scattered Light) "
        "generan patrones de arco en el espectrograma Q-transform que ocupan "
        "un rango de frecuencia estrecho, mientras que los Low Frequency Burst "
        "son estructuralmente mas anchos en frecuencia."
    )
    doc.add_paragraph()

    _add_image(doc, M2_PIPE / "decision_boundaries_2d.png", 5.5,
               "Figura 9. Frontera de decision en el espacio "
               "peak_frequency vs ae_score (otras features fijadas en mediana).")
    doc.add_paragraph()

    _add_image(doc, M2_PIPE / "rule_firing_heatmap.png", 5.5,
               "Figura 10. Heatmap de fuerza de activacion de reglas "
               "(R x N=500 muestras, ordenadas por clase verdadera).")
    doc.add_paragraph()


def build_pipeline(doc):
    _h1(doc, "5. Pipeline End-to-End")

    _h2(doc, "5.1 Diagrama de Flujo")
    _para(doc, "El pipeline completo M1+M2 sigue los siguientes pasos:")
    doc.add_paragraph()

    steps = [
        ("Entrada",        "Ventana cruda 1 s (Q-transform H1/O3a)"),
        ("M1: Normalizar", "Clip + normalizacion global P1=0.025 / P99=7.354"),
        ("M1: Encoder",    "GlitchAE encoder -> vector latente z (dim=32)"),
        ("M1: Decoder",    "Reconstruccion x_hat -> AE-score = MSE(x, x_hat)"),
        ("Deteccion",      "AE-score > umbral -> clasificar; si no -> nominal"),
        ("M2: Features",   "Extraer [peak_freq, ae_score, log_energy, snr, "
                           "duration, bandwidth] desde M1 + metadata CSV"),
        ("M2: Normalizar", "MinMax con rangos almacenados en m2_features.npz"),
        ("M2: ANFIS",      "Fuzzificar -> T-norm -> normalizar -> consecuentes "
                           "LSE -> weighted sum -> softmax -> clase"),
        ("Salida",         "Clase: Loud | Burst | Scatter | Other"),
    ]

    t = doc.add_table(rows=len(steps), cols=2)
    t.style = "Table Grid"
    for i, (step, detail) in enumerate(steps):
        t.rows[i].cells[0].text = step
        t.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        t.rows[i].cells[1].text = detail
        t.rows[i].cells[0].width = Inches(1.8)
        t.rows[i].cells[1].width = Inches(4.5)
    doc.add_paragraph()

    _h2(doc, "5.2 Validacion de Consistencia")
    _para(doc,
        "Se ejecuto el pipeline completo sobre las 500 muestras de run03 y "
        "se comparo la prediccion de cada muestra con la prediccion obtenida "
        "cargando directamente las features precalculadas del archivo "
        "m2_features.npz. El resultado fue:"
    )
    _bullet(doc, "Predicciones iguales: 500/500  (PASS)")
    _bullet(doc, "Tiempo total de ejecucion: 8.7 s (CPU+CUDA)")
    doc.add_paragraph()

    _h2(doc, "5.3 Artefactos Generados")
    _add_table(
        doc,
        headers=["Archivo", "Descripcion"],
        rows=[
            ["m1_v3_outputs/best_m1_ae_v3.pt",          "Checkpoint M1 GlitchAE"],
            ["m1_v3_outputs/normalization_v3.json",      "P1/P99 globales M1"],
            ["m2_data/m2_features.npz",                  "Features + labels 500 muestras"],
            ["m2_outputs/best_m2_anfis.pt",              "Checkpoint M2 ANFIS (standalone)"],
            ["m2_anfis/checkpoints/full_5class_best.pt", "Checkpoint M2 ANFIS (package)"],
            ["m2_outputs/pipeline_eval/*.png",           "4 plots de evaluacion de pipeline"],
        ],
        col_widths=[3.2, 3.1],
    )
    doc.add_paragraph()


def build_limitations(doc):
    _h1(doc, "6. Limitaciones y Trabajo Futuro")

    _h2(doc, "6.1 Limitaciones Actuales")
    _add_table(
        doc,
        headers=["Limitacion", "Impacto", "Mitigacion aplicada"],
        rows=[
            ["Dataset pequeno: 500 muestras evaluacion, 465 entrenamiento M1",
             "Clases minoritarias (Loud=22, Other=32) con F1 bajo",
             "Class weights en loss; balanced split"],
            ["Distribucion desequilibrada: 90 % Burst+Scatter",
             "Other class recall=0.14 en test",
             "Fusion Line+Other para reducir fragmentacion"],
            ["Solo H1/O3a",
             "Generalizacion desconocida a L1 o periodos diferentes",
             "ADR-0016: M1 cerrado hasta ampliar datos"],
            ["5 reglas difusas en espacio 6D",
             "Frontera de decision simplificada",
             "ra=0.15 para clusters naturales sin sobreajuste"],
            ["AE-score como unico feature de M1 para M2",
             "Pierde informacion de la estructura latente z",
             "PCA(z) calculado pero no usado (varianza > 97 % en 2 componentes)"],
        ],
        col_widths=[2.2, 1.8, 2.3],
    )
    doc.add_paragraph()

    _h2(doc, "6.2 Trabajo Futuro")
    _bullet(doc,
        "Ampliar dataset: H1/O3b, L1/O3a, L1/O3b "
        "(objetivo: >5000 muestras etiquetadas)."
    )
    _bullet(doc,
        "M3 -- Regresion auxiliar: correlacion de AE-score con canales "
        "auxiliares (PEM, ASC) para identificar origen instrumental. "
        "Diferido por dependencia de datos de canales auxiliares."
    )
    _bullet(doc,
        "Aumentar reglas ANFIS con mas datos: con >2000 muestras, "
        "ra=0.15 deberia generar 15-25 reglas con mejor cobertura de clases "
        "minoritarias."
    )
    _bullet(doc,
        "Usar embedding latente z completo (32 dims) como feature adicional "
        "para M2, con reduccion UMAP en lugar de PCA."
    )
    _bullet(doc,
        "Despliegue en tiempo real: pipeline M1+M2 en <50 ms por ventana "
        "es factible en GPU (actualmente 8.7 s para 500 muestras = 17 ms/muestra)."
    )


def build_adr(doc):
    _h1(doc, "7. ADR Registry -- Decisiones de Arquitectura")

    _para(doc,
        "Las siguientes Architecture Decision Records (ADR) documentan "
        "las decisiones de diseno mas relevantes del proyecto."
    )
    doc.add_paragraph()

    adrs = [
        {
            "id": "ADR-0015",
            "title": "Normalizacion global P1/P99 sin normalizacion por ventana",
            "status": "Aceptada",
            "context":
                "El autoencoder M1 debe aprender la estructura de ventanas nominales. "
                "Una normalizacion por ventana (P99 local) enmascara la amplitud "
                "absoluta, que es precisamente lo que distingue un glitch intenso "
                "de una ventana nominal.",
            "decision":
                "Calcular P1 y P99 globales sobre el conjunto de entrenamiento "
                "(465 ventanas nominales) y aplicar la misma transformacion "
                "en evaluacion. Almacenar en normalization_v3.json.",
            "consequences":
                "AUROC sube de ~0.70 (sin ADR-0015) a 0.85. "
                "Requiere que el archivo de normalizacion acompane al checkpoint.",
        },
        {
            "id": "ADR-0016",
            "title": "M1 cerrado a AUROC=0.85 sin reentrenamiento",
            "status": "Aceptada",
            "context":
                "El retorno marginal de mejoras en M1 disminuye rapidamente "
                "dado el tamano del dataset. AUROC=0.85 es suficiente para "
                "el proyecto actual.",
            "decision":
                "Congelar el checkpoint M1 (best_m1_ae_v3.pt) y no reentrenar "
                "hasta tener datos de H1/O3b o L1.",
            "consequences":
                "M2 dependera del AE-score de este modelo especifico. "
                "Cualquier cambio en M1 invalida m2_features.npz.",
        },
        {
            "id": "ADR-0017",
            "title": "M2 usa features de M1 encoder + metadata de Gravity Spy",
            "status": "Aceptada",
            "context":
                "El clasificador M2 necesita features discriminativas. "
                "El AE-score de M1 captura la anomalia global, mientras que "
                "peak_frequency, snr, duration y bandwidth del CSV de "
                "Gravity Spy son features fisicas directamente interpretables.",
            "decision":
                "Usar 6 features: [peak_frequency, ae_score, log_energy, snr, "
                "duration, bandwidth]. Descartar los componentes PCA del "
                "vector latente z.",
            "consequences":
                "Interpretabilidad maxima de las reglas ANFIS. "
                "Se descarta informacion del embedding latente (PCA calculado "
                "pero no usado).",
        },
        {
            "id": "ADR-0018",
            "title": "Reduccion de 23 clases a 4 macro-clases",
            "status": "Aceptada",
            "context":
                "Con 500 muestras, muchas clases tienen <5 ejemplos en test. "
                "Un clasificador de 23 clases seria estadisticamente no "
                "evaluable.",
            "decision":
                "Agrupar en 4 macro-clases: Loud, Burst, Scatter, Other. "
                "La clase Line (15 muestras) se fusiona con Other.",
            "consequences":
                "Loss de granularidad para clases Line/Whistle/etc. "
                "Evaluacion estadisticamente valida para Burst y Scatter "
                "(>200 muestras cada una).",
        },
    ]

    for adr in adrs:
        _h2(doc, f"{adr['id']}: {adr['title']}")
        _add_table(
            doc,
            headers=["Campo", "Contenido"],
            rows=[
                ["Estado",       adr["status"]],
                ["Contexto",     adr["context"]],
                ["Decision",     adr["decision"]],
                ["Consecuencias",adr["consequences"]],
            ],
            col_widths=[1.5, 5.0],
        )
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("Building final_report.docx ...")
    doc = Document()

    # Page margins
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    section = doc.sections[0]
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    build_title(doc)
    build_executive_summary(doc)
    doc.add_page_break()

    build_m1(doc)
    doc.add_page_break()

    build_m2(doc)
    doc.add_page_break()

    build_rules(doc)
    doc.add_page_break()

    build_pipeline(doc)
    doc.add_page_break()

    build_limitations(doc)
    doc.add_page_break()

    build_adr(doc)

    out_path = DOCS_DIR / "final_report.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
